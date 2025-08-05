import re
from pathlib import Path

import unicodedata
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .docx_indexer import DocxIndexer


class DocumentProcessor:
    def __init__(self, config, logger):
        self.config = config
        self.logger = logger
        self.url_patterns = [
            (re.compile(transform.from_pattern, re.IGNORECASE), transform.to_pattern)
            for transform in self.config.transform.url_transforms
        ]
        self.logger.extra.update(
            {"location": "", "section": "", "document_name": "", "document_full_path": "", "module": __name__}
        )
        self.current_heading = None

    def _is_in_table(self, paragraph):
        """Check if the paragraph is inside a table cell."""
        if not paragraph or not hasattr(paragraph, "_element"):
            return False

        parent = paragraph._element.getparent()
        depth = 0
        while parent is not None and depth < 10:  # Add depth limit to prevent infinite loops
            self.logger.debug(f"Checking parent tag: {parent.tag}")
            if parent.tag.endswith("tc"):  # tc = table cell
                return True
            parent = parent.getparent()
            depth += 1
        return False

    def _rel_hyperlinks(self, element: Document, doc_index) -> None:
        """Process a section of the document for URL modifications."""
        # Process hyperlinks in relationships
        self.logger.extra.update({"module": "rel_hyperlinks", "task": "rel_URLs"})

        if hasattr(element.part, "rels"):
            for rel_id, rel in element.part.rels.items():
                if rel.reltype == RT.HYPERLINK:
                    original_url = rel.target_ref
                    for pattern, replacement in self.url_patterns:
                        if pattern.search(original_url):
                            para = doc_index.find_paragraph_by_rId(rel_id)
                            self.logger.debug(f"Paragraph:: {para.text}")
                            if para and self._is_in_table(para):
                                self.logger.extra["location"] = "Table"
                                self.logger.extra["match"] = "True"
                                new_url = pattern.sub(replacement, original_url)
                                self.logger.info(f"TABLE: {rel.target_ref} -> {new_url}")
                            else:
                                closest_heading = doc_index.find_closest_heading_above(para)
                                self.logger.extra["location"] = closest_heading if closest_heading else ""
                                self.logger.extra["match"] = "True"
                                new_url = pattern.sub(replacement, original_url)
                                self.logger.info(f"{rel.target_ref} -> {new_url}")

                            self.logger.extra["match"] = "False"
                            rel._target = new_url

    def _para_hyperlinks(self, element: Document, doc_index) -> None:
        self.logger.extra.update({"module": "para_hyperlinks", "task": "para_URLs"})
        for para in element.paragraphs:
            for hyperlink in para.hyperlinks:
                for runs in hyperlink.runs:
                    original_url = runs.text
                    for pattern, replacement in self.url_patterns:
                        if pattern.search(original_url):
                            new_url = pattern.sub(replacement, original_url)
                            closest_heading = doc_index.find_closest_heading_above(para)
                            self.logger.extra["location"] = closest_heading if closest_heading else ""
                            self.logger.extra["match"] = "True"
                            self.logger.info(f"{runs.text} -> {new_url}")
                            self.logger.extra["match"] = "False"
                            runs.text = new_url
                        # Does Not Modify url

    def transform_urls(self, doc: Document, doc_index) -> None:
        """
        Modify URLs in the document.
        There are 2 types or URL's Relationship and Paragraph
        There are 3 Broad locations Headers(multiple), Footers(multiple) and Body
        """
        # Process body text
        self.logger.extra["section"] = "Body"
        self._rel_hyperlinks(doc, doc_index)  # Type A
        self._para_hyperlinks(doc, doc_index)  # Type B

        # Process headers and footers
        for idx, section in enumerate(doc.sections):
            self.logger.extra["section"] = "Header"
            self._rel_hyperlinks(section.header, doc_index)  # Type A
            self._para_hyperlinks(section.header, doc_index)  # Type A

            self.logger.extra["section"] = "Footer"
            self._rel_hyperlinks(section.footer, doc_index)  # Type A

    def transform_styles(self, doc: Document) -> None:
        """Change style names according to configuration."""
        self.logger.extra.update(
            {
                "section": "Whole Document",
                "module": "transform_styles",
            }
        )

        for transform in self.config.transform.style_transforms:
            for style in doc.styles:
                if style.name == transform.from_pattern:
                    style.name = transform.to_pattern
                    self.logger.extra["match"] = "True"
                    self.logger.info(
                        f"Table Style {transform.from_pattern} Found.. Converting, /"
                        f"{transform.from_pattern} → {transform.new_pattern}"
                    )
                    self.logger.extra["match"] = "False"

    def _should_drop_match(self, text):
        """Check if text matches any drop patterns."""
        text = unicodedata.normalize("NFKC", text.strip()).lower()

        for pattern in self.config.transform.drop_matches:
            normalized_pattern = unicodedata.normalize("NFKC", pattern.strip()).lower()

            if normalized_pattern in text:
                self.logger.debug(f"Dropping match for text: '{text[:50]}...'")
                return True
        return False

    def transform_text(self, element: Document, doc_index, transforms):
        """Transform text in document according to configured patterns."""
        self.logger.extra.update(
            {
                "section": "Body",
                "module": "transform_text",
            }
        )

        processed_cells = set()  # Track processed cells by their internal ID

        def process_paragraph(para, cell=None):
            # Skip if this paragraph is in a cell we've already processed
            #        if cell and cell._tc in processed_cells:
            #            return False

            para_text = "".join(run.text for run in para.runs)
            self.logger.debug(f"Paragraph Text: {para.text}")

            if self._should_drop_match(para_text):
                return False

            found_match = False
            for regex in transforms:
                matches = len(re.findall(regex.from_pattern, para_text))
                if matches > 0:
                    if not found_match:
                        trunc_para_text = (para_text[:47] + "...") if len(para_text) > 50 else para_text

                        if para and self._is_in_table(para):
                            self.logger.extra["location"] = "Table"
                        else:
                            closest_heading = doc_index.find_closest_heading_above(para)
                            self.logger.extra["location"] = closest_heading if closest_heading else ""
                        self.logger.extra["match"] = "True"
                        self.logger.info(
                            f"Match: {matches} {'matches' if matches > 1 else 'match'} "
                            f"for {regex.from_pattern}' at paragraph: '{trunc_para_text}'"
                        )
                        self.logger.extra["match"] = "False"
                        self.logger.extra["table_row"] = ""
                        found_match = True

            return found_match

        # Process paragraphs in document body
        for para in element.paragraphs:
            process_paragraph(para)

        # Process paragraphs in tables
        table_no = 0
        for table in element.tables:
            table_no += 1
            row_no = 0
            for row in table.rows:
                row_no += 1
                for cell in row.cells:
                    self.logger.extra["table_row"] = f"{table_no}->{row_no}"
                    # Mark this cell as processed
                    if cell._tc not in processed_cells:
                        processed_cells.add(cell._tc)
                        for para in cell.paragraphs:
                            self.logger.debug(f"In Tables {cell._tc} {para.text} {cell.text}")
                            process_paragraph(para, cell)

        self.logger.extra["table_row"] = ""
        return None

        # TODO: Add Text Transformation

    def process_document(self, input_path: Path, output_path: Path) -> None:
        """Process a single document."""
        self.logger.extra.update({"document_name": input_path.name, "document_full_path": str(input_path.parent)})

        try:
            doc = Document(str(input_path))
            self.logger.extra.update({"section": "NA", "module": "process_document"})

            doc_index = DocxIndexer(doc, self.logger)
            self.logger.debug("-- Index Document --")
            self.logger.debug("-- Start Processing --")

            # Check for non standard Hyperlinks and Log
            # self.logger.debug("Starting Non-Rel-URL Identification")
            # non_rel_hyperlinks(self.logger, input_path)
            # TODO Several of these have to itterate Paragraphs so makes sense to do them in one block
            if self.config.transform.url_transforms:
                self.logger.extra["task"] = "hyperlinks"
                self.logger.debug("Starting URL Identification")
                self.transform_urls(doc, doc_index)

            if self.config.transform.style_transforms:
                self.logger.extra["task"] = "Styles"
                self.logger.debug("Starting Style Identification")
                self.transform_styles(doc)

            if self.config.transform.text_transforms:
                self.logger.extra["task"] = "Text"
                self.logger.debug("Starting Text Identification")
                self.transform_text(doc, doc_index, self.config.transform.text_transforms)

            # Save The Document
            if not self.config.runtime.find_only:
                doc.save(str(output_path))
                self.logger.extra.update({"section": "NA", "task": "Finish", "module": "process_document"})
                self.logger.debug(f"Document saved: {output_path}")

        except Exception as e:
            self.logger.extra["task"] = "ERROR"
            self.logger.error(f"Failed to process {input_path} with error: {str(e).split(':')[0]}")
