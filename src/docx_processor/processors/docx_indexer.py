from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph


class DocxIndexer:
    def __init__(self, doc: Document, logger):
        self.doc = doc
        self.rId_to_paragraph: Dict[str, Paragraph] = {}
        self.paragraph_index: Dict[Paragraph, int] = {}
        self.heading_paragraphs: List[Tuple[Paragraph, int]] = []
        self._build_index()
        self.logging = logger

    def _get_paragraph_id(self, para: Paragraph) -> str:
        """
        Gets the unique paragraph ID from Word's XML structure.
        Falls back to alternative identification if not available.
        """

        # Check if paragraph is None or doesn't have _element
        if para is None or not hasattr(para, "_element") or para._element is None:
            self.logging.warning(f"**Invalid Paragraph Returned: {para}")
            return f"invalid_paragraph_{id(para)}"

        # Try to get the w14:paraId first
        para_id = para._element.get("{http://schemas.microsoft.com/office/word/2010/wordml}paraId")
        if para_id:
            return para_id

        # Fallback to w:id if available
        w_id = para._element.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
        if w_id:
            return w_id

        # Last resort: use combination of properties
        return f"{para.text}|{para.style.name if para.style else ''}|{id(para._element)}"

    def _build_index(self):
        for i, para in enumerate(self.doc.paragraphs):
            para_id = self._get_paragraph_id(para)
            self.paragraph_index[para_id] = i

            if para.style and para.style.name.startswith("Heading"):
                try:
                    heading_level = int(para.style.name[7:])
                except ValueError:
                    heading_level = 0
                self.heading_paragraphs.append((para, heading_level))

            for element in para._element.iter():
                if element.tag.endswith("hyperlink"):
                    rId_value = element.get(f'{{{element.nsmap["r"]}}}id')
                    if rId_value:
                        self.rId_to_paragraph[rId_value] = para
        # Index paragraphs in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Add table paragraphs to the index
                        para_id = self._get_paragraph_id(para)
                        # Use a high index to ensure they come after regular paragraphs
                        self.paragraph_index[para_id] = len(self.doc.paragraphs) + len(self.paragraph_index)

                        for element in para._element.iter():
                            if element.tag.endswith("hyperlink"):
                                rId_value = element.get(f'{{{element.nsmap["r"]}}}id')
                                if rId_value:
                                    self.rId_to_paragraph[rId_value] = para


    def find_paragraph_by_rId(self, rId: str) -> Optional[Paragraph]:
        """
        Retrieves a Paragraph object by its rId from the index.
        """
        return self.rId_to_paragraph.get(rId)

    def find_closest_heading_above(self, paragraph: Paragraph) -> Optional[str]:
        """
        Finds the closest heading paragraph above a given paragraph using the paragraph ID.
        """
        para_id = self._get_paragraph_id(paragraph)
        paragraph_index = self.paragraph_index.get(para_id, -1)

        if paragraph_index == -1:
            return None

        for heading, level in reversed(self.heading_paragraphs):
            heading_id = self._get_paragraph_id(heading)
            if self.paragraph_index.get(heading_id, -1) < paragraph_index:
                return f"H{level} {heading.text}"

        return None
