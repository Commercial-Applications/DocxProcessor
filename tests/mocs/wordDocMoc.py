# Create a new document without TOC or Index
import random

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from faker import Faker

fake = Faker()

final_doc = Document()

TARGET_WORDS = ["FindMe1", "FindMe2", "FindMe3", "FindMe4"]
TARGET_WORDS_COUNT = {word: 20 for word in TARGET_WORDS}
NUM_PAGES = 50
HYPERLINKS_COUNT = 50
MAILTO_COUNT = 50

# Reset counters again
word_inserts = {word: 0 for word in TARGET_WORDS}
links_inserted = 0
mailto_inserted = 0
heading_counters = {1: 0, 2: 0, 3: 0}

mailto_links = [f"mailto:{fake.first_name().lower()}.{fake.last_name().lower()}@testcompany.com" for _ in range(25)]

# Generate fixed links: https://testcompany.com/Test-1 to Test-25
fixed_links = [f"https://testcompany.com/Test-{i}" for i in range(1, 26)]

# Generate 25 random web addresses using Faker
random_links = [fake.url() for _ in range(25)]

fixed_links_remaining = fixed_links.copy()
fixed_link_positions = random.sample(range(NUM_PAGES * 5), 25)  # Pick 25 random positions to insert the links
current_position = 0

heading_counters = {1: 0, 2: 0, 3: 0}


def add_styled_hyperlink(paragraph, url, text):
    """
    Add a styled hyperlink to a paragraph. Style = blue and underlined.
    """
    # Create the relationship ID
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    # Create the <w:hyperlink> element and set the relationship ID
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a <w:r> element
    new_run = OxmlElement("w:r")

    # Create a <w:rPr> element for styling
    rPr = OxmlElement("w:rPr")

    # Style: color = blue
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    # Style: underline = single
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    # Add the hyperlink text
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    # Append run to hyperlink element and hyperlink to paragraph
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def get_numbered_heading(level):
    if level == 1:
        heading_counters[1] += 1
        heading_counters[2] = 0
        heading_counters[3] = 0
        return f"{heading_counters[1]}"
    elif level == 2:
        heading_counters[2] += 1
        heading_counters[3] = 0
        return f"{heading_counters[1]}.{heading_counters[2]}"
    elif level == 3:
        heading_counters[3] += 1
        return f"{heading_counters[1]}.{heading_counters[2]}.{heading_counters[3]}"


for page in range(NUM_PAGES):
    for level in range(1, random.randint(2, 4)):
        heading_number = get_numbered_heading(level)
        heading_text = f"{heading_number} {fake.sentence(nb_words=6)}"
        final_doc.add_heading(heading_text, level=level)
        for _ in range(random.randint(1, 5)):
            paragraph = final_doc.add_paragraph()

            # Add regular text
            text = fake.paragraph(nb_sentences=5)
            words = text.split()
            insert_positions = random.sample(range(len(words)), min(len(words), 2))
            for pos in insert_positions:
                if sum(word_inserts.values()) < 100:
                    choice = random.choice(TARGET_WORDS)
                    if word_inserts[choice] < TARGET_WORDS_COUNT[choice]:
                        words[pos] = choice
                        word_inserts[choice] += 1

            paragraph.add_run(" ".join(words))

            # Insert fixed link if this is one of our randomly chosen positions
            if current_position in fixed_link_positions and fixed_links_remaining:
                link = fixed_links_remaining.pop()
                add_styled_hyperlink(paragraph, link, "Click here")
            # Insert random link if we still have slots available
            elif links_inserted < HYPERLINKS_COUNT - 25:  # Reserve 25 slots for fixed links
                link = random_links[links_inserted]
                add_styled_hyperlink(paragraph, link, "Click here")
                links_inserted += 1

            if mailto_inserted < MAILTO_COUNT and random.random() < 0.05:
                mailto = mailto_links[mailto_inserted]
                add_styled_hyperlink(paragraph, mailto, "Email us")
                mailto_inserted += 1

            current_position += 1

output_path = "./MocWordDoc.docx"
final_doc.save(output_path)
