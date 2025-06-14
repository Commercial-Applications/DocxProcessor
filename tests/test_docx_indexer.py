# %%
from pathlib import Path
from unittest.mock import Mock

import pytest
from docx import Document
from docx.text.paragraph import Paragraph

from docx_processor.processors.docx_indexer import DocxIndexer


@pytest.fixture
def mock_logger():
    logger = Mock()
    logger.logger = Mock()  # Add this for ContextLoggerAdapter compatibility
    return logger


@pytest.fixture
def mock_document():
    doc = Mock(spec=Document)
    doc.paragraphs = []
    return doc


@pytest.fixture
def mock_paragraph():
    para = Mock(spec=Paragraph)
    para._element = Mock()
    para._element.get.return_value = None
    para._element.iter.return_value = []
    para.text = "Test paragraph"
    para.style = Mock()
    para.style.name = "Normal"
    return para


@pytest.fixture
def test_doc_path():
    return Path("tests/data/test_document.docx")


def test_get_paragraph_id_with_para_id(mock_logger, mock_document, mock_paragraph):
    mock_document.paragraphs = []
    indexer = DocxIndexer(mock_document, mock_logger)
    mock_paragraph._element.get.return_value = "test-id-123"

    result = indexer._get_paragraph_id(mock_paragraph)

    assert result == "test-id-123"
    mock_paragraph._element.get.assert_called_with("{http://schemas.microsoft.com/office/word/2010/wordml}paraId")


def test_get_paragraph_id_fallback_to_w_id(mock_logger, mock_document, mock_paragraph):
    mock_document.paragraphs = []
    mock_paragraph._element.get.side_effect = [None, "w-id-456"]
    indexer = DocxIndexer(mock_document, mock_logger)
    result = indexer._get_paragraph_id(mock_paragraph)

    assert result == "w-id-456"
    assert mock_paragraph._element.get.call_count == 2


def test_get_paragraph_id_last_resort(mock_logger, mock_document, mock_paragraph):
    mock_document.paragraphs = []
    mock_paragraph._element.get.return_value = None
    mock_paragraph.style.name = "Heading 1"

    indexer = DocxIndexer(mock_document, mock_logger)
    result = indexer._get_paragraph_id(mock_paragraph)

    assert "Test paragraph|Heading 1|" in result


def test_get_paragraph_id_none_paragraph(mock_logger, mock_document):
    mock_document.paragraphs = []
    indexer = DocxIndexer(mock_document, mock_logger)

    result = indexer._get_paragraph_id(None)

    assert "invalid_paragraph_" in result
    mock_logger.warning.assert_called_once()


def test_build_index(mock_logger, mock_document, mock_paragraph):
    mock_document.paragraphs = [mock_paragraph]
    mock_paragraph._element.get.return_value = "test-id-123"
    mock_paragraph.style.name = "Heading 1"
    mock_paragraph._element.iter.return_value = []
    indexer = DocxIndexer(mock_document, mock_logger)

    assert len(indexer.paragraph_index) == 1
    assert len(indexer.heading_paragraphs) == 1
    assert indexer.heading_paragraphs[0][1] == 1


def test_find_closest_heading_above(mock_logger, mock_document):
    heading = Mock(spec=Paragraph)
    heading._element = Mock()
    heading._element.get.return_value = "heading-id"
    heading._element.iter.return_value = []
    heading.text = "Test Heading"
    heading.style = Mock()
    heading.style.name = "Heading 1"

    content = Mock(spec=Paragraph)
    content._element = Mock()
    content._element.get.return_value = "content-id"
    content._element.iter.return_value = []
    content.style = Mock()
    content.style.name = "Normal"

    mock_document.paragraphs = [heading, content]
    indexer = DocxIndexer(mock_document, mock_logger)

    result = indexer.find_closest_heading_above(content)

    assert result == "H1 Test Heading"


def test_find_paragraph_by_rId(mock_logger, mock_document, mock_paragraph):
    mock_document.paragraphs = [mock_paragraph]

    mock_element = Mock()
    mock_element.tag = "hyperlink"
    mock_element.nsmap = {"r": "test-namespace"}
    mock_element.get.return_value = "test-rid"
    mock_paragraph._element.iter.return_value = [mock_element]

    indexer = DocxIndexer(mock_document, mock_logger)
    result = indexer.find_paragraph_by_rId("test-rid")

    assert result == mock_paragraph
